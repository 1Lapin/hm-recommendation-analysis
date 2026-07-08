# ============================================================
# 生成项目介绍 Word 文档 — H&M个性化推荐策略经营分析
# ============================================================

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ============================================================
# 样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        heading_style.font.size = Pt(18)
        heading_style.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    elif level == 2:
        heading_style.font.size = Pt(14)
        heading_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    elif level == 3:
        heading_style.font.size = Pt(12)
        heading_style.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)

def add_para(text, bold=False, size=None, color=None, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10.5)
    return p

def set_cell_font(cell, text, bold=False, size=Pt(10), color=None, alignment=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = size
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_with_data(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, h, bold=True, size=Pt(10), color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, '1A56DB')
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            set_cell_font(cell, str(val), size=Pt(10))
            if r % 2 == 0:
                set_cell_shading(cell, 'F0F4FF')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacer
    return table

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para('H&M 个性化推荐策略', bold=True, size=Pt(28),
         color=RGBColor(0x1A, 0x56, 0xDB), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
add_para('经营分析项目', bold=True, size=Pt(24),
         color=RGBColor(0x2C, 0x3E, 0x50), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(30))

add_para('—— 基于137万用户 × 10.5万商品 × 3178万条交易数据 ——',
         size=Pt(13), color=RGBColor(0x7F, 0x8C, 0x8D), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(40))

add_para('项目周期：2018-09 ~ 2020-09（27个月数据）', size=Pt(11),
         color=RGBColor(0x7F, 0x8C, 0x8D), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
add_para(f'文档日期：{datetime.now().strftime("%Y年%m月%d日")}', size=Pt(11),
         color=RGBColor(0x7F, 0x8C, 0x8D), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
add_para('关键词：RFM分析 · 用户分群 · 渠道分析 · 品类偏好 · 价格带分析 · A/B测试 · 推荐策略',
         size=Pt(9), color=RGBColor(0x95, 0xA5, 0xA6), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# 目录页（手动）
# ============================================================
doc.add_heading('目  录', level=1)
toc_items = [
    ('一、项目概述', '业务背景、数据规模、核心目标'),
    ('二、技术架构', '数据流、技术栈、项目结构'),
    ('三、数据管道（ETL）', 'Pipeline六步详解、数据质量报告、星型模型设计'),
    ('四、业务诊断', '商品结构诊断、用户数据诊断、核心矛盾定位'),
    ('五、绩效拆解（RFM分析）', '五层RFM分群、GMV驱动因子、月度KPI追踪'),
    ('六、用户分析', '品类偏好建模、价格带锁定、渠道偏好、时间规律、用户画像宽表'),
    ('七、渠道分析', '渠道×品类、渠道×价格、渠道×年龄交叉分析'),
    ('八、推荐策略设计', '5种推荐策略、召回/排序/重排三层架构'),
    ('九、A/B测试框架', '实验设计、评估指标、MAB持续优化'),
    ('十、项目成果与展望', '核心数字摘要、交付物清单、面试追问应答'),
]
for title, desc in toc_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{title}  ')
    run1.bold = True
    run1.font.size = Pt(11)
    run1.font.name = '微软雅黑'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run2 = p.add_run(desc)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    run2.font.name = '微软雅黑'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ============================================================
# 一、项目概述
# ============================================================
doc.add_heading('一、项目概述', level=1)

doc.add_heading('1.1 业务背景', level=2)
add_para('H&M 作为全球快时尚品牌，拥有百万级用户和海量SKU，面临三个核心经营挑战：')

add_table_with_data(
    ['经营挑战', '数据表现', '业务影响'],
    [
        ['商品长尾严重', '14.8%商品购买≤5次，却占库存大头', '库存积压 → 打折清仓 → 利润被蚕食'],
        ['用户留存不稳定', '平均购买间隔大，9000+用户完全无交易', '拉新成本是留存的5倍，流失即亏损'],
        ['推荐缺乏个性', '缺乏推荐系统则所有人看到相同首页', '点击率低 → 转化低 → GMV增长乏力'],
    ],
    [3.5, 4.5, 7.0]
)

add_para('核心假设：如果能为每个用户提供个性化的商品推荐，将显著提升点击率、转化率和人均GMV。')
add_para('核心增长抓手：提升用户复购频次（频次弹性 > 客单价弹性）。')

doc.add_heading('1.2 数据规模', level=2)
add_table_with_data(
    ['数据实体', '记录数', '关键字段', '存储大小'],
    [
        ['用户 (customers)', '1,371,980 行', 'age, club_member_status, fashion_news_frequency', '—'],
        ['商品 (articles)', '105,542 行', 'category_path, colour_group, detail_desc', '—'],
        ['交易 (transactions)', '31,788,324 行', 't_dat, price, sales_channel_id', '4.2 GB'],
        ['时间跨度', '2018-09-20 ~ 2020-09-22', '27个月', '—'],
    ],
    [4.0, 3.0, 5.0, 2.5]
)

doc.add_heading('1.3 核心目标', level=2)
add_para('本项目从0到1构建数据驱动的个性化推荐策略体系，核心目标为：')
add_bullet('提升 GMV（商品交易总额）——北极星指标')
add_bullet('提升用户留存与复购频次')
add_bullet('解决"热门商品垄断、长尾商品曝光不足"的经营矛盾')
add_bullet('输出可复用的数据管道和分析脚本，支撑长期运营决策')

doc.add_heading('1.4 项目整体流程', level=2)
add_para('项目按以下五阶段推进：')
add_table_with_data(
    ['阶段', '核心工作', '产出', '状态'],
    [
        ['Phase 1: 数据处理', 'CSV → 清洗 → 维度建模 → 特征工程', '5张核心表 + 4张Tableau视图', '✅ 完成'],
        ['Phase 2: 经营分析', '业务诊断 → RFM分群 → 用户/渠道分析', '4张分析CSV + 关键数字JSON', '✅ 完成'],
        ['Phase 3: 策略设计', '召回/排序/重排 + 5种推荐策略', '推荐策略设计文档', '✅ 完成'],
        ['Phase 4: 模型训练', 'DSSM双塔 + FAISS + 排序模型', '模型文件 + 离线评估报告', '⏳ 规划中'],
        ['Phase 5: A/B测试', '14天线上实验 + 效果评估', '实验报告 + 运营手册', '⏳ 规划中'],
    ],
    [3.0, 4.5, 4.5, 2.0]
)

doc.add_page_break()

# ============================================================
# 二、技术架构
# ============================================================
doc.add_heading('二、技术架构', level=1)

doc.add_heading('2.1 技术栈', level=2)
add_table_with_data(
    ['层级', '技术', '版本/说明', '用途'],
    [
        ['数据处理', 'Python (Pandas, NumPy)', 'Pandas 2.3, NumPy 1.26', '3300万行ETL + 特征工程'],
        ['可视化', 'Tableau', 'Desktop 2023+', '4张Dashboard（KPI/用户/品类/商品）'],
        ['数据存储', 'CSV文件', '→ 可迁移至MySQL', '原始数据 + 处理后数据'],
        ['版本控制', 'Git + GitHub', '—', '代码和文档管理'],
        ['文档', 'Markdown / Word', '—', '方案设计 + 策略文档 + 分析报告'],
        ['计划引入', 'FAISS + PyTorch', '—', '向量检索引擎 + 双塔模型训练'],
    ],
    [3.0, 4.0, 3.0, 5.5]
)

doc.add_heading('2.2 推荐系统架构（六层）', level=2)
add_para('推荐系统按六层架构设计，数据流自下而上：')

add_table_with_data(
    ['层级', '核心功能', '技术方案', '状态'],
    [
        ['① 数据层', '原始数据清洗 → 维度建模 → 特征宽表', 'Python Pandas 分块读取', '✅ 完成'],
        ['② 特征工程层', '用户特征(14维) + 商品特征(15维)', 'RFM多窗口 + 品类向量', '✅ 完成'],
        ['③ 召回层', '从10万商品中粗筛200个候选', 'DSSM双塔 + FAISS向量检索', '⏳ 规划'],
        ['④ 排序层', '对200候选精排 → Top 50', '多目标排序(CTR+CVR+ATC)', '⏳ 规划'],
        ['⑤ 重排层', '业务规则：库存/新品/多样性', '规则引擎 (品类打散/冷启动加权)', '✅ 已设计'],
        ['⑥ A/B实验层', '50%对照组 vs 50%实验组', '哈希分流 + t检验 + MAB', '✅ 已设计'],
    ],
    [2.5, 4.5, 5.0, 2.0]
)

doc.add_heading('2.3 项目文件结构', level=2)
add_para('项目采用清晰的模块化结构，代码和文档分离：')

files_struct = [
    ['01_data_pipeline.py', '全流程数据管道（6步）', '~37分钟', '5张主表 + 4张Tableau视图'],
    ['02_business_analysis.py', '经营分析脚本', '~8分钟', '4张CSV + 关键数字JSON'],
    ['项目方案_H&M推荐系统.md', '总体技术方案文档', '—', '6层架构 + 阶段规划'],
    ['推荐策略设计文档_基于规则.md', '5种推荐策略 + A/B方案', '—', 'SQL逻辑 + Excel模板'],
    ['数据分析工作思考过程.md', '数据分析方法论', '—', '思维框架 + 协作边界'],
    ['processed/', '处理后数据（不入库）', '9个CSV', 'dim_*, fact_*, feature_*, tableau_*'],
    ['business_analysis/', '经营分析结果', '5个文件', 'RFM GMV/渠道品类/年龄渠道/KPI'],
]
add_table_with_data(
    ['文件/目录', '说明', '耗时', '产出'],
    files_struct,
    [4.5, 4.5, 2.0, 5.0]
)

doc.add_page_break()

# ============================================================
# 三、数据管道
# ============================================================
doc.add_heading('三、数据管道（ETL）', level=1)

doc.add_heading('3.1 Pipeline 六步流程', level=2)
add_para('01_data_pipeline.py 完整实现从原始CSV到特征宽表的全流程，共6步：')

pipeline_steps = [
    ['Step 1', '加载原始数据', '读取3个CSV（articles/customers/transactions）', 'articles: 105,542行\ncustomers: 1,371,980行\ntransactions: 31,788,324行'],
    ['Step 2', '清洗 articles', '填充detail_desc空值 + 构建品类路径 + 计算商品购买次数 + 热度分层', 'dim_article: 105,542行\n含category_path / popularity_tier'],
    ['Step 3', '清洗 customers', '★核心修复：用交易行为反推Active字段\n（66%为空但99.3%有交易→修正活跃用户46万→136万）\n+ 年龄段分组 + 会员状态标准化', 'dim_customer: 1,371,980行\n含age_group / is_active'],
    ['Step 4', '处理 transactions', '时间特征提取 + 引用完整性验证 + 训练/验证集按时间划分（最后7天=验证集）', 'fact_transaction: 31,788,324行\n含data_split标签'],
    ['Step 5', 'RFM用户特征', '多窗口频率(7d/30d/90d) + 品类多样性 + 衍生指标(recency/online_ratio/间隔) + 五层RFM分群', 'feature_user_rfm: 1,356,132行\n20个特征字段'],
    ['Step 6', '商品协同特征', '时间趋势(up/stable/down) + 回购率 + 渠道偏好 + 购买者画像(年龄均值/多样性)', 'feature_article_stats: 104,229行\n15个特征字段'],
]
add_table_with_data(
    ['步骤', '任务', '关键操作', '产出'],
    pipeline_steps,
    [1.5, 2.5, 6.5, 5.0]
)

doc.add_heading('3.2 关键技术决策', level=2)
add_bullet('分块读取（Chunked Reading）：fact_transaction 达 4.2GB，采用 chunksize=500,000 分块处理，避免内存溢出')
add_bullet('Active 字段修复：发现66%用户 Active 为空但99.3%有实际交易——用交易行为反推修正，这是数据清洗中最重要的逻辑修正')
add_bullet('训练/验证集划分：按时间切分（最后7天为验证集），模拟真实线上推荐场景（用历史数据预测未来行为）')
add_bullet('RFM分群阈值：基于实际数据分布设定，非机械等分——高价值活跃(recency≤30天 & frequency≥27次)、活跃(recency≤90天 & frequency≥9次)等')

doc.add_heading('3.3 星型数据模型', level=2)
add_para('数据模型采用星型架构，以 fact_transaction 为事实表，dim_article/dim_customer 为维度表，feature_user_rfm 和 feature_article_stats 为特征宽表：')
add_para('''
  fact_transaction (31,788,324行)
  ├── customer_id → dim_customer (1,371,980行)
  │                 └── age_group, is_active, club_member_status
  ├── article_id  → dim_article (105,542行)
  │                 └── category_path, popularity_tier, colour_group
  ├── 用户维度特征 → feature_user_rfm (1,356,132行, 20维)
  │                 └── rfm_segment, recency_days, online_ratio, unique_departments
  └── 商品维度特征 → feature_article_stats (104,229行, 15维)
                    └── trend_direction, repurchase_rate, avg_buyer_age
''')

doc.add_heading('3.4 数据质量关键发现', level=2)
add_table_with_data(
    ['问题', '严重程度', '影响', '处理方式'],
    [
        ['66%用户Active为空，但99.3%有交易', '🔴 严重', '不能用Active做特征', '用交易行为反推，修正活跃用户46万→136万'],
        ['14.8%商品购买≤5次', '🟡 中等', '召回层冷启动问题', '重排层冷启动加权15% + 热门兜底'],
        ['fashion_news出现"None"/"NONE"', '🟡 中等', '特征编码会分裂', '统一标准化为"NONE"'],
        ['detail_desc 416行空值', '🟢 轻微', '文本特征缺失', '用产品信息(prod_name+product_type+department)填充'],
        ['9,699用户无任何交易记录', '🟢 轻微', '新用户冷启动', '热门兜底策略'],
    ],
    [4.0, 2.0, 3.5, 5.5]
)

doc.add_page_break()

# ============================================================
# 四、业务诊断
# ============================================================
doc.add_heading('四、业务诊断', level=1)

doc.add_heading('4.1 商品结构诊断', level=2)
add_para('基于10.5万SKU的购买分布分析，发现严重的"热门垄断、长尾消失"问题：')

add_table_with_data(
    ['指标', '数值', '解读'],
    [
        ['总SKU数', '105,542', '—'],
        ['购买≤5次的商品', '15,664个 (14.8%)', '长尾积压——这些商品占用库存但几乎卖不动'],
        ['其中冷启动(≤3次)', '大量商品', '新品上市后缺乏曝光机会，陷入"越没曝光越卖不动"的死循环'],
        ['热门商品(≥1000次)', '7,810个 (7.4%)', '贡献了53.5%的销量——极少数商品垄断了大部分曝光和销售'],
        ['结论', '—', '热门集中 vs 长尾消失 → "首页千篇一律、新品永无出头之日"'],
    ],
    [5.0, 4.0, 7.0]
)

doc.add_heading('4.2 用户数据诊断', level=2)
add_para('发现一个关键的数据质量问题：Active 字段与交易行为严重背离。')
add_table_with_data(
    ['维度', '修复前', '修复后'],
    [
        ['Active = "1.0" 的用户', '约46万', '—'],
        ['Active 为空的用户', '约91万 (66%)', '—'],
        ['实际有交易记录的用户', '—', '约136万 (99.3%)'],
        ['修正后的活跃用户', '—', '约136万 (从46万修正)'],
    ],
    [5.0, 4.5, 4.5]
)
add_para('★ 关键发现：如果直接用 Active 字段做特征，会把90万活跃用户误判为"非活跃"——这是一个会导致推荐策略严重偏差的数据质量问题。修复后活跃用户从46万→136万，数据与行为事实一致。')

doc.add_heading('4.3 核心矛盾定位', level=2)
add_para('综合商品结构和用户数据，定位到H&M当前经营的核心矛盾：')
add_bullet('矛盾一：热门垄断 —— 7.4%热门商品贡献53.5%销量，导致首页展示面狭窄')
add_bullet('矛盾二：长尾消失 —— 14.8%商品几乎无曝光机会，库存周转效率极低')
add_bullet('矛盾三：标签失真 —— 66%用户活跃状态被误判，导致运营无法精准触达')
add_para('这三个矛盾相互强化：热门越垄断 → 长尾越没曝光 → 用户看到的越单一 → 数据越集中在热门 → 循环恶化。推荐系统的核心价值就在于打破这个循环。')

doc.add_page_break()

# ============================================================
# 五、绩效拆解
# ============================================================
doc.add_heading('五、绩效拆解（RFM分析）', level=1)

doc.add_heading('5.1 五层RFM分群', level=2)
add_para('基于 Recency（最近购买时间）、Frequency（购买频次）、Monetary（消费金额）三个维度，将136万用户划分为5类价值客群：')

add_table_with_data(
    ['RFM分群', '用户数', '占比', 'GMV贡献', '平均频次', '购买间隔中位'],
    [
        ['高价值活跃', '142,309', '10.5%', '约39%', '83.6次', '3.0天'],
        ['活跃用户', '279,220', '20.6%', '中等', '17.3次', '12.8天'],
        ['潜在流失', '317,557', '23.4%', '中等偏低', '5.3次', '40.2天'],
        ['已流失', '254,366', '18.8%', '很低', '1.9次', '180+天'],
        ['一般用户', '362,680', '26.7%', '较低', '3.2次', '70.0天'],
    ],
    [2.5, 2.0, 1.5, 1.8, 2.0, 2.5]
)

add_para('★ 核心发现：仅10.5%的高价值用户贡献了约39%的GMV，他们是平台的核心资产。同时23.4%的潜在流失用户是"最值得投入唤醒资源"的群体。')

doc.add_heading('5.2 GMV驱动因子拆解', level=2)
add_para('将GMV拆解为 GMV = 用户数 × 人均频次 × 客单价，对比各分群：')
add_bullet('频次方差极大：高价值活跃用户均频83.6次 vs 已流失用户均频1.9次，相差44倍')
add_bullet('客单价方差极小：各分群间客单价差异远小于频次差异')
add_bullet('★ 结论：频次弹性 >> 客单价弹性 → 提升用户复购频次是核心增长抓手，比"提价"更有效且风险更低')
add_bullet('购买间隔分布：P25=14天、P50=34天、P75=76天、P90=190天——提升空间巨大')

doc.add_heading('5.3 月度KPI追踪', level=2)
add_para('建立25个月趋势基线（交易量/活跃用户/GMV/客单价/线上占比/周末占比），识别季节性规律：')
add_bullet('旺季峰值：12月（假日季促销）均交易量最高')
add_bullet('低谷期：1月～2月交易量明显下降（节后消费疲软）')
add_bullet('线上占比：稳定在29.6%左右，无明显季节波动')
add_bullet('周末占比：稳定在28.6%，验证了"周末逛街"的消费习惯')

doc.add_heading('5.4 策略定调', level=2)
add_para('基于以上分析，确定策略优先级：')
add_bullet('第一优先级：唤醒23.4%潜在流失用户 —— 存量最大、唤醒成本最低')
add_bullet('第二优先级：拉升20.6%活跃用户的购买频次 —— 增长空间最大')
add_bullet('第三优先级：稳住10.5%高价值用户 —— 他们是GMV基本盘')
add_bullet('核心逻辑：唤醒+拉升的潜在GMV提升空间 >> 通过提价带来的GMV增长，且风险更低')

doc.add_page_break()

# ============================================================
# 六、用户分析
# ============================================================
doc.add_heading('六、用户分析', level=1)

doc.add_heading('6.1 品类偏好建模', level=2)
add_para('核心创新：构建 preference_score = user_ratio / global_ratio 指标，消除"热门品类误导"。')
add_para('举例说明：如果用户A买了100件商品，其中10件是Ladieswear（占比10%），而全球Ladieswear占比为50%，则 preference_score = 10%/50% = 0.2 → 实际上用户A并不偏好Ladieswear。但如果只看绝对购买次数（10件），会误以为"买了10件，应该喜欢吧"。')
add_para('这个指标确保了推荐逻辑的准确性——不是"推最热门的"，而是"推用户真正偏好的"。')

doc.add_heading('6.2 价格带锁定', level=2)
add_para('★ 这是本项目最核心的用户洞察之一：')
add_para('通过对10万抽样用户的全量交易分析，发现 99.5% 的用户 ≥80% 的购买集中在单一价格带。')
add_table_with_data(
    ['指标', '数值', '含义'],
    [
        ['≥80%购买在单一价格带', '99.5% 用户', '价格偏好极度集中'],
        ['≥90%购买在单一价格带', '更高比例', '用户价格忠诚度极高'],
        ['CV < 0.3（购买≥3次用户）', '高比例', '价格变异系数低，证明偏好稳定'],
    ],
    [5.0, 3.0, 6.0]
)
add_para('业务启示：价格匹配在推荐中极其关键。如果给用户推了他不习惯的价位商品，几乎一定不会购买。价格偏好一旦形成极为稳定，推荐时必须在用户的价格带内。')

doc.add_heading('6.3 渠道偏好识别', level=2)
add_para('基于 online_ratio 将用户分为三类渠道偏好：')
add_table_with_data(
    ['渠道偏好类型', '用户占比', '定义', '推荐策略启示'],
    [
        ['线上为主', '26.1%', 'online_ratio > 70%', '优先推线上热销款 + 基础款'],
        ['线下为主', '60.0%', 'online_ratio < 30%', '优先推新品 + 趋势款 + 高客单价'],
        ['混合渠道', '14.0%', '30% ≤ online_ratio ≤ 70%', '根据当前渠道实时切换策略'],
    ],
    [3.0, 2.0, 4.0, 5.0]
)

doc.add_heading('6.4 时间消费规律', level=2)
add_para('分析全量交易的天维度分布：')
add_bullet('周末购买占比：28.6%（显著高于均分14.3%的预期）')
add_bullet('工作日高峰：周一和周五购买稍高（上班前/周末前准备）')
add_bullet('工作日分布较为均匀（12.8%～15.3%每日）')
add_para('业务启示：推荐推送的最佳时机为周五下午+周末上午；周末适合推"新品发现型"推荐，工作日适合"高效补充型"推荐。')

doc.add_heading('6.5 用户画像宽表', level=2)
add_para('最终构建14维用户特征宽表，直接支撑运营分群触达：')
add_para('画像字段：age_group（年龄段）、preference_score（品类偏好度）、price_band（价格带）、online_ratio（渠道偏好）、recency_days（最近购买天数）、frequency_total（总购买频次）、rfm_segment（RFM分群）、unique_departments（品类多样性）、avg_purchase_interval（平均购买间隔）、weekend_ratio（周末偏好）等。')

doc.add_page_break()

# ============================================================
# 七、渠道分析
# ============================================================
doc.add_heading('七、渠道分析', level=1)

doc.add_heading('7.1 渠道×品类交叉', level=2)
add_para('通过对3178万条交易按销售渠道（线上29.6% / 线下70.4%）与品类的交叉分析：')
add_table_with_data(
    ['品类', '线上占比 vs 全局', '差异', '解读'],
    [
        ['Jewellery', '线上显著高于线下', '+1.6pp', '饰品标准化程度高，适合线上购买'],
        ['Knitwear', '线上偏高', '+1.4pp', '基础款无需试穿'],
        ['Hair Accessories', '线上偏好', '+pp', '小件配饰，线上决策成本低'],
        ['Swimwear', '线下显著偏高', '-5.2pp', '泳装需要试穿体验'],
        ['Jersey', '线下偏高', '-1.7pp', '运动面料需要手感体验'],
    ],
    [3.0, 3.5, 2.0, 5.5]
)

doc.add_heading('7.2 渠道×价格交叉', level=2)
add_para('计算线上与线下客单价对比（全量3178万条交易）：')
add_bullet('线上客单价 / 线下客单价 = 0.77')
add_bullet('线上用户倾向于购买价格更低的商品，符合"线上比价、线下体验"的消费行为逻辑')
add_bullet('线下渠道天然适合推高客单价商品和新品——用户在门店更愿意为"体验"买单')

doc.add_heading('7.3 渠道×年龄交叉', level=2)
add_para('分析不同年龄段的线上购物偏好差异：')
add_bullet('25-34岁 线上占比最低（约25.2%）——这个群体最倾向于线下购物')
add_bullet('65岁以上 线上占比最高（约36.6%）——可能因为行动不便更依赖线上')
add_bullet('年龄段与渠道偏好存在显著差异，分年龄段的渠道触达策略应有区别')

doc.add_heading('7.4 渠道策略落地', level=2)
add_para('基于以上三项交叉分析，输出分渠道推荐方案：')
add_table_with_data(
    ['渠道', '推荐策略', '核心逻辑'],
    [
        ['线上', '高性价比基础款 + 高复购品类', '线上用户价格敏感、偏好标准化商品'],
        ['线下', '当季新品 + 趋势款 + 高客单价商品', '线下用户愿为体验付费，适合新品首发'],
        ['混合型用户', '根据 online_ratio 动态调整权重', '用历史渠道偏好预测当前最相关推荐'],
    ],
    [2.5, 5.5, 6.0]
)

doc.add_page_break()

# ============================================================
# 八、推荐策略设计
# ============================================================
doc.add_heading('八、推荐策略设计', level=1)

doc.add_heading('8.1 策略体系总览', level=2)
add_para('设计5种推荐策略，按用户状态分层使用，形成完整的推荐策略矩阵：')

add_table_with_data(
    ['策略编号', '策略名称', '适用条件', '推荐位占比', '核心逻辑', '可调参数'],
    [
        ['0', '热门兜底', '新用户/无购买', '100%（兜底）', '全局销量降序Top20', '时间窗口'],
        ['1', '混合探索', '购买≤2次', '60%热门+40%随机', '热门为主+探索性推荐', '热门/随机比例'],
        ['2', '品类偏好', '购买≥3次', '40%（8位）', '基于preference_score推荐偏好品类下热销商品', 'TopN品类数、每品类取几个'],
        ['3', '价格带', '购买≥3次', '25%（5位）', '匹配用户价格带±1档的热销商品', '价格箱数量、浮动档位'],
        ['4', '协同过滤', '购买≥5次', '25%（5位）', '基于Jaccard共购相似度', '最近购买取几条、相似度阈值'],
        ['5', '颜色偏好', '购买≥3次', '10%（2位）', '匹配用户偏好颜色组+同品类热销', '取前几个颜色'],
    ],
    [1.5, 2.0, 2.0, 2.0, 4.5, 3.0]
)

doc.add_heading('8.2 三层推荐架构', level=2)

doc.add_heading('召回层（粗筛：10万→200）', level=3)
add_para('目标：从10万商品中快速筛出200个候选，保证"该召回的都召回"。')
add_bullet('双塔模型（DSSM）：用户塔 + 商品塔分别编码为128维向量 → FAISS向量检索')
add_bullet('品类规则召回补充：用户历史偏好品类下热销商品强制进入候选集 → 解决长尾召回不足')
add_bullet('多样性召回通道：随机采样不同department的热门商品混入候选集')

doc.add_heading('排序层（精排：200→50）', level=3)
add_para('目标：对200个候选按用户偏好打分排序。')
add_bullet('多目标排序：CTR（点击率）+ CVR（购买率）+ ATC（加购率）')
add_bullet('加权融合：final_score = 0.3 × p_ctr + 0.5 × p_cvr + 0.2 × p_atc')
add_bullet('权重设定来自敏感性分析——CVR权重最高因为直接指向GMV')

doc.add_heading('重排层（业务规则：50→20）', level=3)
add_para('目标：加入业务约束，确保推荐"既准又赚钱"。')
add_table_with_data(
    ['规则', '具体操作', '目的'],
    [
        ['库存过滤', '库存<阈值的商品降权或移除', '避免推荐买不到的商品'],
        ['新品加权', '冷启动商品(购买≤3次)加权15%', '给新品曝光机会，均衡库存'],
        ['多样性打散', '同品类(department)最多出现3个', '避免推荐列表全是同类商品'],
        ['价格多样性', '至少保留2个价格带的商品', '覆盖不同预算需求'],
        ['去已购', '过去7天已购买的商品降权', '避免重复推荐'],
    ],
    [3.0, 5.5, 5.5]
)

doc.add_heading('8.3 偏好度计算创新', level=2)
add_para('品类偏好度（preference_score）计算方法：')
add_bullet('Step 1：计算用户在各品类的购买占比 (user_ratio)')
add_bullet('Step 2：计算该品类的全局占比 (global_ratio)')
add_bullet('Step 3：preference_score = user_ratio / global_ratio')
add_bullet('> 1.0 = 比平均人群更偏好该品类 → 推荐信号')
add_bullet('> 2.0 = 显著偏好 → 强推荐信号')
add_para('这个设计的关键洞察：不是买得多就是喜欢。"相对于全局的偏好"才是做推荐时该用的真实信号。')

doc.add_page_break()

# ============================================================
# 九、A/B测试框架
# ============================================================
doc.add_heading('九、A/B测试框架', level=1)

doc.add_heading('9.1 实验设计', level=2)
add_table_with_data(
    ['实验要素', '设计方案'],
    [
        ['实验周期', '14天（覆盖两个完整周，消除周末效应 + 新推荐的新奇效应）'],
        ['分流方式', 'customer_id 哈希取模 → hash(id) % 2 == 0 → 对照组A / == 1 → 实验组B'],
        ['对照组A (50%流量)', '热门商品兜底推荐（按全局销量降序）'],
        ['实验组B (50%流量)', '完整推荐系统（品类偏好+价格带+协同过滤+重排）'],
        ['AA验证', '上线前先做AA测试（两组用相同策略），确认分流均匀无偏'],
        ['数据剔除', '前2天数据标记为"新奇效应期"，分析时做敏感性检验'],
    ],
    [3.5, 12.5]
)

doc.add_heading('9.2 评估指标体系', level=2)
add_table_with_data(
    ['层级', '指标', '计算公式', '优先级'],
    [
        ['北极星指标', '人均GMV', 'SUM(price) / COUNT(DISTINCT user)', '⭐⭐⭐'],
        ['核心指标', 'CTR（点击率）', '点击次数 / 曝光次数', '⭐⭐⭐'],
        ['核心指标', 'CVR（购买转化率）', '购买次数 / 曝光次数', '⭐⭐⭐'],
        ['观察指标', '推荐覆盖率', '被推荐过的商品数 / 总商品数', '⭐⭐'],
        ['观察指标', '品类多样性', '推荐列表中不同department的数量', '⭐⭐'],
        ['观察指标', '冷启动商品曝光占比', '冷启动商品曝光 / 总曝光', '⭐⭐'],
        ['诊断指标', '客单价', 'GMV / 购买订单数', '⭐'],
    ],
    [2.5, 4.0, 5.5, 2.0]
)

doc.add_heading('9.3 显著性检验与判断标准', level=2)
add_para('采用双样本异方差 t检验（双尾），显著性水平 α = 0.05。因同时观察多个指标，对次要指标使用 Bonferroni 校正。')
add_table_with_data(
    ['场景', '判断', '行动'],
    [
        ['p < 0.05, 提升 > 5%', '✅ 策略有效', '全量上线该策略'],
        ['p < 0.05, 提升 0~5%', '⚠️ 微弱有效', '优化策略后再测'],
        ['p < 0.05, 提升 < 0', '❌ 策略有害', '立即下线，排查原因'],
        ['p ≥ 0.05', '❓ 无法判断', '延长实验或加大样本量'],
    ],
    [4.0, 3.0, 7.0]
)

doc.add_heading('9.4 MAB持续优化', level=2)
add_para('长期方案：用 Thompson Sampling（多臂老虎机算法）自动调整流量分配。')
add_bullet('每个策略 = 一个"臂"；臂的表现 = Beta分布(购买次数+1, 未购买次数+1)')
add_bullet('每次推荐时从每个臂的Beta分布采样 → 选采样值最大的策略')
add_bullet('效果：好策略自动获得更多流量，差策略自动减少，无需人工干预')

doc.add_heading('9.5 预期效果', level=2)
add_table_with_data(
    ['指标', '预期提升', '原因'],
    [
        ['CTR（点击率）', '+10~20%', '品类偏好让推荐"更对胃口"'],
        ['CVR（购买转化率）', '+8~15%', '价格带匹配降低"看了买不起"'],
        ['推荐覆盖率', '+100~300%', '从只推热门→推长尾商品'],
        ['冷启动商品曝光', '从0→5%', '重排层有意识地给新品机会'],
        ['人均GMV', '+5~10%', 'CVR提升 + 价格带匹配的综合效果'],
    ],
    [4.0, 3.0, 7.0]
)

doc.add_page_break()

# ============================================================
# 十、项目成果与展望
# ============================================================
doc.add_heading('十、项目成果与展望', level=1)

doc.add_heading('10.1 核心数字摘要', level=2)
add_para('以下是本项目的关键分析发现（所有数字由 Python 脚本实际运行产出，可复现验证）：')

summary_items = [
    ['数据规模', '137万用户 × 10.5万商品 × 3178万交易记录（4.2GB）'],
    ['时间跨度', '27个月（2018-09-20 ~ 2020-09-22）'],
    ['商品长尾', '14.8%商品购买≤5次，7.4%热门商品贡献53.5%销量'],
    ['数据修正', '66%用户Active为空但99.3%有交易 → 修正活跃用户46万→136万'],
    ['RFM分群', '5层：高价值活跃10.5% / 活跃20.6% / 潜在流失23.4% / 已流失18.8% / 一般26.7%'],
    ['GMV集中度', '仅10.5%高价值用户贡献约39% GMV'],
    ['增长抓手', '频次弹性 >> 客单价弹性 → 提升复购是核心'],
    ['价格带', '99.5%用户≥80%购买集中在单一价格带'],
    ['渠道分布', '线上29.6% vs 线下70.4%；用户：线上为主26.1%/线下为主60.0%/混合14.0%'],
    ['线上vs线下价差', '线上客单价为线下的0.77倍'],
    ['周末效应', '周末购买占比28.6%，高于均分预期'],
    ['年龄×渠道', '25-34岁线上占比最低(25.2%)，65岁以上最高(36.6%)'],
]
add_table_with_data(
    ['分析维度', '核心发现'],
    summary_items,
    [3.5, 12.5]
)

doc.add_heading('10.2 代码与文档交付物', level=2)
add_table_with_data(
    ['交付物', '形式', '规模/耗时', '用途'],
    [
        ['01_data_pipeline.py', 'Python脚本', '~400行, ~37分钟', '全流程数据管道，一键重现所有数据处理'],
        ['02_business_analysis.py', 'Python脚本', '~600行, ~8分钟', '经营分析，产出所有关键数字'],
        ['项目方案_H&M推荐系统.md', 'Markdown文档', '~300行', '总体技术方案（6层架构+阶段规划）'],
        ['推荐策略设计文档_基于规则.md', 'Markdown文档', '~550行', '5种推荐策略+SQL逻辑+A/B方案'],
        ['数据分析工作思考过程.md', 'Markdown文档', '~400行', '数据分析方法论+思维框架+协作边界'],
        ['processed/ (9个CSV)', '处理后数据', '合计约4.8GB', '5张主表+4张Tableau视图'],
        ['business_analysis/ (5个文件)', '分析结果', '4CSV+1JSON', 'RFM GMV/渠道品类/年龄渠道/KPI'],
        ['Tableau Dashboard ×4', '可视化看板', '4张交互式看板', 'KPI/用户分群/品类销售/商品趋势'],
    ],
    [5.0, 2.5, 3.5, 5.0]
)

doc.add_heading('10.3 项目亮点总结', level=2)
add_bullet('数据规模大：处理3300万行、4.2GB数据，使用分块读取策略避免内存溢出')
add_bullet('分析有深度：从业务诊断→绩效拆解→用户分析→渠道分析，完整的经营分析链路')
add_bullet('方法论扎实：preference_score消除热门品类误导、99.5%价格带锁定等创新分析方法')
add_bullet('落地可执行：5种推荐策略 + A/B实验框架 + 重排规则，直接可交付算法/运营团队')
add_bullet('工程化思维：所有代码可复用、数字可复现，数据更新后一键重跑')
add_bullet('决策有依据：频次弹性>客单价弹性——用数据证明"提升复购"比"提价"更有效')

doc.add_heading('10.4 面试追问应答', level=2)

add_para('Q1: "你诊断出热门垄断、长尾消失的矛盾，怎么解决的？"', bold=True)
add_para('设计了"三层开窗策略"：① 召回层品类偏好召回（40%推荐位），让用户收到真正偏好品类下的商品而非纯热门；② 重排层冷启动加权15% + 品类打散（同品类≤3个），强制给新品和长尾商品曝光；③ 长期用MAB自动调整各策略流量分配。核心思想：不消灭热门推荐（它本身有效），而是在热门基础上"开窗"给长尾和新品。')

add_para('Q2: "怎么确定提升复购比提价更有效？"', bold=True)
add_para('两个角度交叉验证：① RFM数据——用户购买频次方差极大（P25=3次 vs P90=59次），提升空间大；而客单价受价格带锚定效应影响，99.5%用户购买集中在单一价格带，强行提价可能导致用户流失。② 业务逻辑——快时尚本质是"高频低单价"模式，频繁上新驱动复购是天生的增长引擎，提价违背品牌定位。')

add_para('Q3: "渠道分析对推荐系统有什么具体影响？"', bold=True)
add_para('落地为三个具体规则：① 如果能区分用户渠道 → 实时切换推荐策略（线上推平价基础款、线下推新品趋势款）；② 如果无法区分 → 根据用户历史online_ratio偏好做推荐；③ 新品推广时优先推给线下偏好用户（新品需要体验，线下渠道天然适合首发）。')

add_para('Q4: "项目最终落地了什么？"', bold=True)
add_para('三样东西：① 两个可复用的Python脚本（数据管道+经营分析），数据更新后一键重跑，所有数字自动验证；② 一份推荐策略设计文档（5种策略+重排规则+A/B方案），算法和产品可以直接实现；③ 4个Tableau看板，运营团队日常就能看KPI趋势和用户分群，不需要找数据分析师取数。')

doc.add_heading('10.5 后续规划', level=2)
add_table_with_data(
    ['阶段', '工作内容', '预计耗时', '交付物'],
    [
        ['短期', 'Tableau 4张看板制作', '一个下午', '交互式可视化看板'],
        ['中期', 'DSSM双塔模型训练 + FAISS索引构建', '1-2天', '模型文件 + 离线评估'],
        ['中期', '排序模型训练（多目标）', '1-2天', 'LightGBM/XGBoost排序模型'],
        ['长期', '线上A/B实验 + MAB持续优化', '14天实验 + 持续', '实验报告 + 运营手册'],
    ],
    [2.0, 5.0, 3.0, 5.0]
)

# ============================================================
# 尾页
# ============================================================
doc.add_page_break()
doc.add_heading('附录：技术环境', level=1)
add_table_with_data(
    ['项目', '说明'],
    [
        ['Python版本', 'Python 3.8+'],
        ['核心库', 'Pandas 2.3, NumPy 1.26'],
        ['运行环境', 'Windows / macOS / Linux'],
        ['数据来源', 'Kaggle - H&M Personalized Fashion Recommendations'],
        ['GitHub仓库', 'https://github.com/1Lapin/hm-recommendation-analysis'],
        ['开发工具', 'VS Code / Jupyter Notebook'],
    ],
    [4.0, 12.0]
)

# ============================================================
# 保存
# ============================================================
output_path = os.path.join(OUTPUT_DIR, 'H&M个性化推荐策略_经营分析项目介绍.docx')
doc.save(output_path)
print(f'✅ Word 文档已生成: {output_path}')
print(f'   文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')
