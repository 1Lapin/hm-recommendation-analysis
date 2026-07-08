# ============================================================
# 将 18 张图表插入 Word 文档 — 按六个阶段对应插入
# ============================================================
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

CHARTS_DIR = 'charts'
WORD_IN = 'H&M个性化推荐策略_经营分析项目介绍.docx'
WORD_OUT = 'H&M个性化推荐策略_经营分析项目介绍_含图表.docx'

# Chart mapping: (phase_section_keyword, chart_filename, caption)
CHART_MAP = [
    # Phase 1: 业务诊断
    ('四、业务诊断', 'phase1/01_商品购买次数分布.png', '图1: 商品购买次数分布 — 14.8%商品≤5次 vs 7.4%热门贡献53.5%销量'),
    ('四、业务诊断', 'phase1/02_商品销量贡献结构.png', '图2: 商品销量贡献结构 — 商品数量占比与销量贡献不均衡'),
    ('四、业务诊断', 'phase1/03_Active修复前后对比.png', '图3: Active字段修复前后对比 — 46万→136万 (+195%)'),
    # Phase 2: 绩效拆解
    ('五、绩效拆解', 'phase2/04_RFM五层分群.png', '图4: RFM五层分群双轴图 — 10.5%高价值用户贡献39% GMV'),
    ('五、绩效拆解', 'phase2/05_购买频次对比.png', '图5: 各分群购买频次对比 — 高频83.6次 vs 流失1.9次 (44倍差距)'),
    ('五、绩效拆解', 'phase2/06_月度KPI趋势.png', '图6: 月度KPI三面板趋势 — 25个月GMV/活跃用户/客单价'),
    # Phase 3: 用户分析
    ('六、用户分析', 'phase3/07_渠道偏好分布.png', '图7: 用户渠道偏好分布 — 线下为主60.0%/线上为主26.1%/混合14.0%'),
    ('六、用户分析', 'phase3/08_价格带集中度.png', '图8: 价格带集中度仪表盘 — 99.5%用户≥80%购买集中在单一价格带'),
    ('六、用户分析', 'phase3/09_时间消费规律.png', '图9: 周度消费规律 — 周末购买占比28.6%,显著高于均匀预期'),
    # Phase 4: 渠道分析
    ('七、渠道分析', 'phase4/10_渠道品类交叉.png', '图10: 渠道×品类交叉 — Jewellery线上偏好 vs Swimwear线下偏好'),
    ('七、渠道分析', 'phase4/11_渠道年龄交叉.png', '图11: 渠道×年龄交叉 — 25-34岁线上最低(25.2%),65岁以上最高(36.6%)'),
    ('七、渠道分析', 'phase4/12_渠道价格对比.png', '图12: 渠道×价格对比 — 线上客单价仅为线下的0.77倍'),
    # Phase 5: 推荐策略设计
    ('八、推荐策略设计', 'phase5/13_推荐位分配.png', '图13: 推荐位分配 — 品类偏好40%/价格带25%/协同过滤25%/颜色10%'),
    ('八、推荐策略设计', 'phase5/14_推荐架构漏斗.png', '图14: 推荐系统架构漏斗 — 10万→200→50→20的四层过滤'),
    ('八、推荐策略设计', 'phase5/15_偏好度计算创新.png', '图15: 偏好度计算创新 — preference_score消除"热门品类误导"'),
    # Phase 6: A/B测试框架
    ('九、A/B测试框架', 'phase6/16_预期效果对比.png', '图16: A/B实验预期效果对比 — CTR+15%/CVR+10%/GMV+8%'),
    ('九、A/B测试框架', 'phase6/17_MAB流量优化.png', '图17: MAB流量优化模拟 — 好策略自动获得更多流量'),
    ('九、A/B测试框架', 'phase6/18_显著性判断矩阵.png', '图18: 显著性判断标准 — p值+提升率四维决策矩阵'),
]

print('>>> 加载 Word 文档...')
doc = Document(WORD_IN)

# 设置样式
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_chart_after_section(doc, section_keyword, chart_path, caption):
    """在找到 section_keyword 的 heading 后，跳到该 section 的结尾插入图表"""
    if not os.path.exists(chart_path):
        print(f'  ⚠️ Not found: {chart_path}')
        return False

    # Find the section heading paragraph
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading') and section_keyword in p.text:
            target_idx = i
            break

    if target_idx is None:
        print(f'  ⚠️ Section "{section_keyword}" not found')
        return False

    # Find the next Heading 1 or end of document
    insert_idx = len(doc.paragraphs) - 1
    for i in range(target_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name == 'Heading 1':
            insert_idx = i - 1
            break

    # Insert chart before the next section (at insert_idx position)
    # If there's a page break at insert_idx, insert before it

    # Strategy: insert a new paragraph with chart after the insert_idx paragraph
    # Find the paragraph element at insert_idx
    ref_para = doc.paragraphs[insert_idx]

    # Add caption
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_para.add_run(caption)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6b, 0x7c, 0x93)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Move caption paragraph after ref_para
    ref_para._element.addnext(cap_para._element)

    # Add image
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(chart_path, width=Inches(5.8))

    # Move image after caption
    cap_para._element.addnext(img_para._element)

    # Spacer
    spacer = doc.add_paragraph()
    cap_para._element.addnext(spacer._element)
    img_para._element.addnext(spacer._element)

    return True

print('>>> 插入图表到对应章节...')
inserted = 0
for section_kw, chart_rel, caption in CHART_MAP:
    chart_path = os.path.join(CHARTS_DIR, chart_rel)
    if add_chart_after_section(doc, section_kw, chart_path, caption):
        inserted += 1
        print(f'  ✅ {chart_rel} → "{section_kw}"')

# Save
doc.save(WORD_OUT)
print(f'\n✅ 已插入 {inserted}/{len(CHART_MAP)} 张图表')
print(f'📄 输出: {WORD_OUT}')
print(f'📏 文件大小: {os.path.getsize(WORD_OUT)/1024:.0f} KB')
